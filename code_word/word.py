from docx import Document
from docx.shared import Inches
import os

# 创建Word文档对象
doc = Document()

# 添加标题
doc.add_heading('期末大作业报告', 0)

# 创建图片保存文件夹
if not os.path.exists("images"):
    os.mkdir("images")

# 1. 概要
doc.add_heading('1. 概要', level=1)
doc.add_paragraph(
    """随着数据分析和机器学习技术的发展，数据的可视化成为了数据分析中不可或缺的一部分。尤其在二手车市场分析中，通过对车辆数据进行深入分析，能够揭示出市场的潜在规律及趋势，从而为决策者提供有价值的参考。本报告基于二手车数据集，利用 Python 中的 Pandas、Matplotlib 和 Seaborn 等可视化工具，对数据进行了清洗、处理和可视化，展示了二手车市场中的一些重要特征和趋势。

    通过对数据的初步分析，首先筛选出未上牌车辆的比例并对其进行统计；接着，通过对车辆的品牌、车龄、里程和价格等特征进行进一步分析，我们生成了多种图表，展示了不同品牌的销售占比、车龄分布、车型价格分布以及价格与里程的关系。我们还对奥迪品牌进行了详细的分析，展示了其不同车型的里程变化情况，进一步探讨了奥迪品牌与价格的关系。

    在整个数据处理过程中，我们进行了数据清洗和转换，确保数据格式的统一性和可操作性。最终，基于这些分析结果，我们生成了八个关键图表，并将这些图表插入到 Word 报告中，使得报告更具可读性和可操作性。本报告为二手车市场的分析提供了一个基本框架，并展示了如何使用 Python 实现数据清洗、分析和可视化。"""
)

# 2. 相关技术
doc.add_heading('2. 相关技术', level=1)
doc.add_paragraph(
    """在本项目中，我们主要使用了以下几种技术和工具：

    - **Pandas**：作为 Python 中用于数据处理的核心库，Pandas 提供了强大的数据结构（如 DataFrame）和数据处理功能（如数据清洗、处理缺失值、数据转换等）。在本项目中，Pandas 主要用于读取 CSV 格式的二手车数据集，并对数据进行筛选、清洗、转换和聚合等处理。

    - **Matplotlib**：Matplotlib 是 Python 中最常用的绘图库之一，能够生成高质量的静态、动态和交互式图表。在本项目中，Matplotlib 用于生成各种类型的图表，包括饼图、条形图、散点图、折线图、箱线图和热力图等。

    - **Seaborn**：Seaborn 是建立在 Matplotlib 基础上的高级可视化库，提供了更为简洁和优美的接口，适用于数据科学中的统计图表绘制。Seaborn 提供了便捷的小提琴图、热力图等功能，帮助我们进行更精细的分析和展示。

    - **python-docx**：为了自动化报告的生成，我们使用了 `python-docx` 库，它是一个可以用来创建和修改 Word 文档的 Python 库。通过 `python-docx`，我们可以将数据分析结果、图表及对应的说明插入到报告中，并最终生成 Word 格式的报告文档。

    在数据清洗和转换阶段，Pandas 用于处理和转换数据中的字符串和数值，例如从 `Boarding_time` 列中提取出年份，并将价格和里程字段中的单位（如 "万"）去除，转化为数值类型。之后，使用 Matplotlib 和 Seaborn 生成图表，这些图表展示了二手车数据的各个方面，包括品牌占比、车龄、价格分布等信息。"""
)

# 3. 具体实现
doc.add_heading('3. 具体实现', level=1)
doc.add_paragraph(
    """以下是实现过程的详细代码，包括生成的8个图表以及对应的代码。所有图表将保存在 'images' 文件夹中，以便插入到最终的 Word 报告中。

    1. 数据预处理：首先，我们从 CSV 文件中读取二手车数据，并进行初步的清洗与转换，例如从 `Boarding_time` 列中提取年份信息，去除价格与里程字段中的单位等。

    2. 生成图表：使用 Matplotlib 和 Seaborn 创建了八个重要的图表，展示了二手车市场的各个方面：
        - 车辆销售品牌占比（饼图）
        - 车龄分布（条形图）
        - 奥迪车型的里程与价格关系（散点图）
        - 部分奥迪车型里程变化情况（折线图）
        - 奥迪品牌车型新车价格分布（小提琴图）
        - 雷达图
        - 价格分布情况（箱线图）
        - 月度销售热度（热力图）
        
    下面是相关代码和图表：
    
    前提：python基础环境配置完成，pip 下载以下依赖。
    
    python脚本输入：
    import pandas as pd
    import numpy as np
    import matplotlib.pyplot as plt
    from matplotlib import rcParams
    import seaborn as sns
    """
)

# 图表 1：车辆销售品牌占比（饼图）
# 插入饼图
doc.add_paragraph("图 1: 车辆销售品牌占比")
doc.add_picture("images/i1.png", width=Inches(5))

# 插入对应代码
doc.add_paragraph("代码：")
code1 = '''
# 确保 'car' 是一个已定义的 DataFrame，并且包含 'Brand' 列
# car = pd.read_csv('your_file.csv')  # 示例加载数据的方式
# 设置字体以支持中文显示
rcParams['font.family'] = 'sans-serif'
rcParams['font.sans-serif'] = ['SimHei']
# 修改此处，使用推荐的方式获取品牌销量前十的数据
car_top10 = pd.Series(car["Brand"]).value_counts().sort_values(ascending=False)[:10]
# 使用更简洁的方法创建 brand 和 brand_num 列表
brand = list(car_top10.index) + ["其他"]
brand_num = list(car_top10.values) + [car.shape[0] - sum(car_top10.values)]
# 将品牌和数量组合成字典并转换为 Series 对象
data = dict(zip(brand, brand_num))
df = pd.Series(data)
# 绘制饼图
plt.figure(figsize=(10, 10))
df.plot(kind="pie", title="车辆销售品牌占比", autopct="%.2f%%")
plt.ylabel('')  # 移除默认的 y 标签
plt.show()
'''
doc.add_paragraph(code1)

# 图表 2：车龄分布（条形图）
# 插入条形图
doc.add_paragraph("图 2: 车龄分布")
doc.add_picture("images/i2.png", width=Inches(5))

# 插入对应代码
doc.add_paragraph("代码：")
code2 = '''
car_year = car['year'].value_counts()
car_year = car_year.to_frame().sort_index()
car_year.plot(figsize=(12, 8), kind='bar', title="车龄")
plt.show()
'''
doc.add_paragraph(code2)

# 图表 3：奥迪车型的里程与价格关系（散点图）
# 插入散点图
doc.add_paragraph("图 3: 奥迪车型的里程与价格关系")
doc.add_picture("images/i3.png", width=Inches(5))

# 插入对应代码
doc.add_paragraph("代码：")
code3 = '''
index = car['Brand'].isin(['奥迪'])
some_cars = car.loc[index, :]
some_cars
some_cars.plot.scatter(
    figsize=(12, 8),
    x='Km', y='Sec_price',
    s=20, marker='o',
    alpha=0.9,
    linewidths=0.3,
    edgecolors='k'
)
plt.show()
'''
doc.add_paragraph(code3)


# 图表 4：部分奥迪车型里程变化情况
doc.add_paragraph("图 4: 部分奥迪车型里程变化情况")
doc.add_picture("images/i4.png", width=Inches(5))

# 插入对应代码
doc.add_paragraph("代码：")
code4 = '''
# 读取数据
data = pd.read_csv("./../datas/second_cars_info.csv", encoding="gbk")
# 使用matplotlib绘制不同车型的里程（Km）的折线图（多条折线，每个车型一条，示例中选取部分车型示意，可根据需求调整）
# 选取部分车型进行示例展示，你可以根据实际需求修改这里选取的车型
# 数据清洗
selected_models = ['奥迪A6L 2006款 2.4 CVT 舒适型', '奥迪A8L 2013款 45 TFSI quattro舒适型', '奥迪Q5 2013款 40TFSI 舒适型']
subset_data = data[data['Name'].isin(selected_models)]

# 绘制折线图
for model in selected_models:
    model_data = subset_data[subset_data['Name'] == model]
    plt.plot(model_data['Km'], label=model)

plt.xlabel('车辆序号（示例）')
plt.ylabel('里程（万公里）')
plt.title('部分奥迪车型里程变化情况')
plt.legend()
plt.show()
'''
doc.add_paragraph(code4)

# 图表 5：奥迪品牌车型新车价格分布（小提琴图）
doc.add_paragraph("图 5: 奥迪品牌车型新车价格分布（小提琴图）")
doc.add_picture("images/i5.png", width=Inches(5))

# 插入对应代码
doc.add_paragraph("代码：")
code5 = '''
# 使用seaborn绘制不同品牌（这里都是奥迪，可拓展为多品牌对比情况）与价格（New_price）的关系小提琴图（示例中可看到价格分布情况）
# 绘制小提琴图，这里只有奥迪品牌，可拓展为多品牌对比
sns.violinplot(x='Brand', y='New_price', data=data)
plt.xlabel('品牌')
plt.ylabel('新车价格（万）')
plt.title('奥迪品牌车型新车价格分布（小提琴图）')
plt.show()
'''
doc.add_paragraph(code5)

# 图表 6：雷达图
# 插入散点图
doc.add_paragraph("图 6: 雷达图")
doc.add_picture("images/i6.png", width=Inches(5))

# 插入对应代码
doc.add_paragraph("代码：")
code6 = '''
labels = ["A", "B", "C", "D"]
num_vars = len(labels)

values = [3, 2, 4, 5]  # 示例数据
values += values[:1]  # 闭合曲线

angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
angles += angles[:1]  # 闭合曲线

fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
ax.fill(angles, values, color="red", alpha=0.25)
ax.plot(angles, values, color="red", linewidth=2)
ax.set_yticklabels([])
ax.set_xticks(angles[:-1])
ax.set_xticklabels(labels)
plt.show()
'''
doc.add_paragraph(code6)

# 图表 7：# 箱线图
# 插入散点图
doc.add_paragraph("图 7: # 箱线图")
doc.add_picture("images/i7.png", width=Inches(5))

# 插入对应代码
doc.add_paragraph("代码：")
code7 = '''
data = [2, 3, 4, 5, 6, 7, 8, 9, 10]  # 示例数据
plt.boxplot(data)
plt.title("价格分布情况")
plt.ylabel("价格")
plt.show()
'''
doc.add_paragraph(code7)


# 图表 8：热力图（需要安装seaborn库）
# 插入散点图
doc.add_paragraph("图 8: 热力图（需要安装seaborn库）")
doc.add_picture("images/i8.png", width=Inches(5))

# 插入对应代码
doc.add_paragraph("代码：")
code8 = '''
data = [[1, 2, 3], [4, 5, 6], [7, 8, 9]]  # 示例数据
df = pd.DataFrame(data, columns=['Jan', 'Feb', 'Mar'])
heatmap_data = df.T
sns.heatmap(heatmap_data, annot=True, cmap="YlGnBu")
plt.title("月度销售热度")
plt.show()
'''
doc.add_paragraph(code8)

# 4. 致谢
doc.add_heading('4. 致谢', level=1)
doc.add_paragraph(
    """首先，感谢提供二手车数据集，使得本次分析能够顺利进行。感谢 Python 开发社区，特别是 Pandas、Matplotlib 和 Seaborn 等库的开发者，提供了强大的数据分析和可视化支持。尤其是这些工具能够让我们快速地完成数据清洗、处理和可视化任务，提高了工作效率。感谢导师在本项目中的指导与支持，在数据清洗、分析和可视化的过程中提供了宝贵的意见和建议，帮助我克服了遇到的困难和挑战。此外，感谢同学们在项目讨论中的积极参与，他们提出了许多有价值的想法和建议，使得我对二手车市场有了更深入的理解和认识。最后，感谢家人的理解与支持，他们在我忙碌的过程中给予了我很多鼓励和帮助，使我能够顺利完成此项工作。"""
)

# 保存文档
doc.save("期末大作业报告.docx")

print("报告已成功保存为'期末大作业报告.docx'")